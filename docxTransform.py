from docx import Document
import codecs

doc = Document(
    "D:\LearningPYTHON\Курсач\Кр_Лекция_1_Знакомство_со_средой.docx")
res_edit = []
res = []
tab_symbol = "<font style='color:white'>----</font>"

for paragraph in doc.paragraphs:
    res_str = ""
    for word in paragraph.runs:
        if word.font.color.rgb != None:
            res_str += f"<font style='color:#{word.font.color.rgb}'>{word.text.replace('    ',tab_symbol)}</font>"
        else:
            res_str += word.text
    res_edit.append(res_str)

for paragraph in doc.paragraphs:
    res.append(paragraph.text)
    res.append(paragraph.style.name)
    res.append(paragraph.style.font.size)
    res.append("\n")



with codecs.open('D:\LearningPYTHON\Курсач\Кр_Лекция_1_Знакомство_со_средой.txt', 'w', "utf-8") as f:
    for line in res_edit:
        f.write(f"{line}\n")

with codecs.open('D:\LearningPYTHON\Курсач\Кр111_Лекция_1_Знакомство_со_средой.txt', 'w', "utf-8") as f:
    for line in res:
        f.write(f"{line}\n")